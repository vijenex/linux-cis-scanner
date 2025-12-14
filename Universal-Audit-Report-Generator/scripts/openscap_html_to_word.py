#!/usr/bin/env python3
"""
OpenSCAP HTML to Word Report Generator
Parses OpenSCAP HTML scan results and generates a comprehensive CIS audit report in Word format
"""

import os
import sys
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from collections import defaultdict

def set_cell_color(cell, color):
    """Set background color for table cell"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)

def parse_html_report(html_file):
    """Parse OpenSCAP HTML report and extract key information"""
    with open(html_file, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')
    
    # Extract system information
    system_info = {}
    
    # Get target system name
    target_elem = soup.find('td', string='Evaluation target')
    if target_elem:
        system_info['target'] = target_elem.find_next_sibling('td').get_text(strip=True)
    
    # Get IP addresses
    ip_addresses = []
    ip_section = soup.find('h4', string='Addresses')
    if ip_section:
        ip_list = ip_section.find_next('ul')
        if ip_list:
            for li in ip_list.find_all('li'):
                ip_text = li.get_text(strip=True)
                # Extract IPv4 addresses
                if 'IPv4' in ip_text:
                    ip_match = re.search(r'(\d+\.\d+\.\d+\.\d+)', ip_text)
                    if ip_match and ip_match.group(1) != '127.0.0.1':
                        ip_addresses.append(ip_match.group(1))
    
    system_info['ip_addresses'] = ip_addresses
    
    # Get scan dates
    started_elem = soup.find('th', string='Started at')
    if started_elem:
        system_info['started'] = started_elem.find_next_sibling('td').get_text(strip=True)
    
    finished_elem = soup.find('th', string='Finished at')
    if finished_elem:
        system_info['finished'] = finished_elem.find_next_sibling('td').get_text(strip=True)
    
    # Parse rule results
    rules = []
    rule_rows = soup.find_all('tr', class_='rule-overview-leaf')
    
    for row in rule_rows:
        rule = {}
        
        # Get rule ID
        rule_id = row.get('id', '')
        rule['id'] = rule_id
        
        # Get rule title
        title_link = row.find('a')
        if title_link:
            rule['title'] = title_link.get_text(strip=True)
        
        # Get severity
        severity_cell = row.find('td', class_='rule-severity')
        if severity_cell:
            rule['severity'] = severity_cell.get_text(strip=True)
        
        # Get result
        result_cell = row.find('td', class_='rule-result')
        if result_cell:
            result_div = result_cell.find('div')
            if result_div:
                abbr = result_div.find('abbr')
                if abbr:
                    rule['result'] = abbr.get_text(strip=True)
        
        # Get CIS reference from data-references attribute
        data_refs = row.get('data-references', '')
        if data_refs:
            try:
                import json
                refs = json.loads(data_refs)
                if 'cis' in refs:
                    rule['cis_id'] = refs['cis'][0] if refs['cis'] else ''
            except:
                pass
        
        if rule.get('title'):
            rules.append(rule)
    
    return system_info, rules

def generate_word_report(scan_results, output_file):
    """Generate comprehensive Word document from scan results"""
    
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('RHEL 8 CIS Benchmark Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('Consolidated Compliance Assessment', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    
    total_systems = len(scan_results)
    total_rules = sum(len(result['rules']) for result in scan_results)
    
    doc.add_paragraph(f'This report presents the results of the CIS (Center for Internet Security) benchmark audit conducted on {total_systems} RHEL 8 system(s).')
    doc.add_paragraph(f'Audit Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph()
    
    # System Inventory
    doc.add_heading('System Inventory', 1)
    
    # Create systems table
    sys_table = doc.add_table(rows=1, cols=4)
    sys_table.style = 'Table Grid'
    
    hdr_cells = sys_table.rows[0].cells
    hdr_cells[0].text = 'System Name'
    hdr_cells[1].text = 'IP Address'
    hdr_cells[2].text = 'Scan Start'
    hdr_cells[3].text = 'Scan End'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_color(cell, 'D3D3D3')
    
    for result in scan_results:
        row_cells = sys_table.add_row().cells
        row_cells[0].text = result['system_info'].get('target', 'N/A')
        row_cells[1].text = ', '.join(result['system_info'].get('ip_addresses', ['N/A']))
        row_cells[2].text = result['system_info'].get('started', 'N/A')
        row_cells[3].text = result['system_info'].get('finished', 'N/A')
    
    doc.add_page_break()
    
    # Compliance Summary
    doc.add_heading('Compliance Summary', 1)
    
    # Aggregate results by environment
    prod_systems = []
    nonprod_systems = []
    
    for result in scan_results:
        # Determine environment based on folder structure or naming
        if 'Prod' in result.get('source_path', ''):
            prod_systems.append(result)
        else:
            nonprod_systems.append(result)
    
    # Count pass/fail by severity
    severity_counts = {'high': {'pass': 0, 'fail': 0}, 
                      'medium': {'pass': 0, 'fail': 0}, 
                      'low': {'pass': 0, 'fail': 0}}
    
    for result in scan_results:
        for rule in result['rules']:
            severity = rule.get('severity', 'unknown').lower()
            result_status = rule.get('result', '').lower()
            
            if severity in severity_counts:
                if result_status == 'pass':
                    severity_counts[severity]['pass'] += 1
                elif result_status == 'fail':
                    severity_counts[severity]['fail'] += 1
    
    doc.add_paragraph(f"High Severity: {severity_counts['high']['pass']} passed, {severity_counts['high']['fail']} failed")
    doc.add_paragraph(f"Medium Severity: {severity_counts['medium']['pass']} passed, {severity_counts['medium']['fail']} failed")
    doc.add_paragraph(f"Low Severity: {severity_counts['low']['pass']} passed, {severity_counts['low']['fail']} failed")
    
    doc.add_page_break()
    
    # Failed Controls Detail
    doc.add_heading('Failed Controls Detail', 1)
    
    # Group failed rules by CIS ID
    failed_by_cis = defaultdict(lambda: {'prod': [], 'nonprod': [], 'rule_info': None})
    
    for result in scan_results:
        env = 'prod' if 'Prod' in result.get('source_path', '') else 'nonprod'
        ip = ', '.join(result['system_info'].get('ip_addresses', ['N/A']))
        
        for rule in result['rules']:
            if rule.get('result', '').lower() == 'fail':
                cis_id = rule.get('cis_id', 'Unknown')
                failed_by_cis[cis_id][env].append(ip)
                if not failed_by_cis[cis_id]['rule_info']:
                    failed_by_cis[cis_id]['rule_info'] = rule
    
    # Sort by CIS ID
    sorted_cis_ids = sorted(failed_by_cis.keys(), key=lambda x: [int(n) if n.isdigit() else n for n in re.split(r'(\d+)', x)])
    
    for cis_id in sorted_cis_ids:
        data = failed_by_cis[cis_id]
        rule = data['rule_info']
        
        # Add control heading
        doc.add_heading(f"CIS Control {cis_id}: {rule['title']}", 2)
        
        # Add severity
        severity = rule.get('severity', 'unknown')
        p = doc.add_paragraph()
        p.add_run('Severity: ').bold = True
        p.add_run(severity.upper())
        
        # Add description
        p = doc.add_paragraph()
        p.add_run('Description: ').bold = True
        p.add_run(f"This control ensures {rule['title'].lower()}.")
        
        # Add affected systems table
        doc.add_paragraph().add_run('Affected Systems:').bold = True
        
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        # Headers
        table.rows[0].cells[0].text = 'Passed (Prod)'
        table.rows[0].cells[1].text = 'Passed (Non-Prod)'
        table.rows[0].cells[2].text = 'Failed (Prod)'
        table.rows[0].cells[3].text = 'Failed (Non-Prod)'
        
        # Color headers
        set_cell_color(table.rows[0].cells[0], 'd4edda')
        set_cell_color(table.rows[0].cells[1], 'd4edda')
        set_cell_color(table.rows[0].cells[2], 'f8d7da')
        set_cell_color(table.rows[0].cells[3], 'f8d7da')
        
        for cell in table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Data
        table.rows[1].cells[0].text = 'None'
        table.rows[1].cells[1].text = 'None'
        table.rows[1].cells[2].text = '\n'.join(data['prod']) if data['prod'] else 'None'
        table.rows[1].cells[3].text = '\n'.join(data['nonprod']) if data['nonprod'] else 'None'
        
        # Color data cells
        if data['prod']:
            set_cell_color(table.rows[1].cells[2], 'f8d7da')
        if data['nonprod']:
            set_cell_color(table.rows[1].cells[3], 'f8d7da')
        
        doc.add_paragraph()
    
    # Save document
    doc.save(output_file)
    print(f'✅ Report generated: {output_file}')

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 openscap_html_to_word.py <path_to_rhel8_folder>")
        print("Example: python3 openscap_html_to_word.py /Users/satish.korra/Downloads/RHEL-8")
        sys.exit(1)
    
    base_path = sys.argv[1]
    
    if not os.path.exists(base_path):
        print(f"Error: Path does not exist: {base_path}")
        sys.exit(1)
    
    print(f"Scanning for OpenSCAP HTML reports in: {base_path}")
    
    # Find all HTML files
    scan_results = []
    
    for root, dirs, files in os.walk(base_path):
        for file in files:
            if file.endswith('.html'):
                html_path = os.path.join(root, file)
                print(f"Processing: {html_path}")
                
                try:
                    system_info, rules = parse_html_report(html_path)
                    scan_results.append({
                        'source_path': html_path,
                        'system_info': system_info,
                        'rules': rules
                    })
                except Exception as e:
                    print(f"Error processing {html_path}: {e}")
    
    if not scan_results:
        print("No HTML reports found!")
        sys.exit(1)
    
    print(f"\nFound {len(scan_results)} scan result(s)")
    
    # Generate output filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = os.path.join(base_path, f"RHEL8-CIS-Audit-Report-{timestamp}.docx")
    
    # Generate report
    print("\nGenerating Word document...")
    generate_word_report(scan_results, output_file)
    
    print(f"\n✅ Complete! Report saved to: {output_file}")

if __name__ == "__main__":
    main()
