#!/usr/bin/env python3
"""
RHEL 8 CIS Audit Report Generator
Generates comprehensive audit reports for RHEL 8 systems with proper CIS control parsing
"""
import os
import re
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from striprtf.striprtf import rtf_to_text

def set_cell_color(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)

def parse_rtf_file(rtf_path):
    """Parse RTF file and extract text content"""
    if not os.path.exists(rtf_path):
        print(f"Warning: RTF file not found: {rtf_path}")
        return ""
    
    try:
        with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
            rtf_content = f.read()
        return rtf_to_text(rtf_content)
    except Exception as e:
        print(f"Error parsing RTF file {rtf_path}: {e}")
        return ""

def extract_failed_controls(text_content):
    """Extract failed controls from OpenSCAP text output"""
    failed_controls = {}
    lines = text_content.split('\n')
    
    current_control = None
    for line in lines:
        line = line.strip()
        
        # Look for control IDs and titles
        if 'fail' in line.lower() and any(char.isdigit() for char in line):
            # Extract control information
            parts = line.split()
            for i, part in enumerate(parts):
                if re.match(r'^\d+\.\d+', part):
                    control_id = part
                    # Get title from remaining parts
                    title_parts = parts[i+1:]
                    title = ' '.join([p for p in title_parts if p.lower() != 'fail' and p.lower() != 'medium' and p.lower() != 'high' and p.lower() != 'low'])
                    
                    failed_controls[control_id] = {
                        'title': title,
                        'status': 'FAIL',
                        'severity': 'medium'  # Default severity
                    }
                    break
    
    return failed_controls

def get_cis_control_details(control_id):
    """Get detailed CIS control information"""
    # Comprehensive CIS RHEL 8 control database
    cis_controls = {
        "1.1.1.1": {
            "title": "Ensure mounting of cramfs filesystems is disabled",
            "description": "The cramfs filesystem type is a compressed read-only Linux filesystem embedded in small footprint systems. A cramfs image can be used without having to first decompress the image.",
            "remediation": "Edit or create a file in the /etc/modprobe.d/ directory ending in .conf and add the following line: install cramfs /bin/true. Run the following command to unload the cramfs module: rmmod cramfs",
            "impact": "Disabling unused filesystem modules reduces attack surface."
        },
        "1.1.1.2": {
            "title": "Ensure mounting of freevxfs filesystems is disabled",
            "description": "The freevxfs filesystem type is a free version of the Veritas type filesystem. This is the primary filesystem type for HP-UX operating systems.",
            "remediation": "Edit or create a file in the /etc/modprobe.d/ directory ending in .conf and add the following line: install freevxfs /bin/true. Run the following command to unload the freevxfs module: rmmod freevxfs",
            "impact": "Removing support for unneeded filesystem types reduces the local attack surface."
        },
        "1.1.2.1": {
            "title": "Ensure /tmp is configured",
            "description": "The /tmp directory is a world-writable directory used for temporary storage by all users and some applications.",
            "remediation": "Configure /tmp as a separate partition or tmpfs. Edit /etc/fstab and add the appropriate entry for /tmp.",
            "impact": "Files saved to /tmp will be lost when system is rebooted if using tmpfs."
        },
        "1.2.1": {
            "title": "Ensure package manager repositories are configured",
            "description": "Systems need to have package manager repositories configured to ensure they can receive software updates.",
            "remediation": "Configure your package manager repositories according to site policy.",
            "impact": "Proper repository configuration is essential for system updates."
        },
        "1.3.1": {
            "title": "Ensure SELinux is installed",
            "description": "SELinux provides Mandatory Access Control.",
            "remediation": "Run the following command to install SELinux: dnf install libselinux",
            "impact": "SELinux provides additional security layer through mandatory access controls."
        },
        "2.1.1": {
            "title": "Ensure xinetd is not installed",
            "description": "The eXtended InterNET Daemon (xinetd) is an open source super daemon that replaced the original inetd daemon.",
            "remediation": "Run the following command to remove xinetd: dnf remove xinetd",
            "impact": "Removing xinetd reduces attack surface."
        },
        "3.1.1": {
            "title": "Disable unused network protocols and devices",
            "description": "The Linux kernel modules support several network protocols that are not commonly used.",
            "remediation": "Disable unused network protocols by blacklisting kernel modules.",
            "impact": "Reducing available network protocols decreases attack surface."
        },
        "4.1.1": {
            "title": "Ensure auditing is enabled",
            "description": "Turn on the auditd daemon to record system events.",
            "remediation": "Run the following commands: systemctl --now enable auditd",
            "impact": "Auditing provides accountability and forensic capabilities."
        },
        "5.1.1": {
            "title": "Ensure cron daemon is enabled",
            "description": "The cron daemon is used to execute batch jobs on the system.",
            "remediation": "Run the following command to enable cron: systemctl --now enable crond",
            "impact": "Cron is required for scheduled system tasks."
        },
        "5.2.1": {
            "title": "Ensure SSH Protocol is configured",
            "description": "SSH server configuration should follow security best practices.",
            "remediation": "Configure SSH server settings in /etc/ssh/sshd_config according to security requirements.",
            "impact": "SSH configuration changes may affect remote access capabilities."
        }
    }
    
    return cis_controls.get(control_id, {
        "title": f"Control {control_id}",
        "description": "Description not available in database.",
        "remediation": "Refer to official CIS RHEL 8 Benchmark documentation.",
        "impact": "Impact assessment required."
    })

def collect_server_results(base_path):
    """Collect audit results from all servers"""
    results = {
        'prod_servers': [],
        'uat_servers': [],
        'all_controls': {}
    }
    
    # Process Prod servers
    prod_path = os.path.join(base_path, 'Prod')
    if os.path.exists(prod_path):
        for server_file in os.listdir(prod_path):
            if server_file.endswith('.txt.rtf'):
                server_ip = server_file.replace('.txt.rtf', '')
                results['prod_servers'].append(server_ip)
                
                rtf_path = os.path.join(prod_path, server_file)
                text_content = parse_rtf_file(rtf_path)
                failed_controls = extract_failed_controls(text_content)
                
                for control_id, control_data in failed_controls.items():
                    if control_id not in results['all_controls']:
                        results['all_controls'][control_id] = {
                            'title': control_data['title'],
                            'failed_prod': set(),
                            'failed_uat': set(),
                            'passed_prod': set(),
                            'passed_uat': set()
                        }
                    results['all_controls'][control_id]['failed_prod'].add(server_ip)
    
    # Process UAT servers
    uat_path = os.path.join(base_path, 'UAT')
    if os.path.exists(uat_path):
        for server_file in os.listdir(uat_path):
            if server_file.endswith('.txt.rtf'):
                server_ip = server_file.replace('.txt.rtf', '')
                results['uat_servers'].append(server_ip)
                
                rtf_path = os.path.join(uat_path, server_file)
                text_content = parse_rtf_file(rtf_path)
                failed_controls = extract_failed_controls(text_content)
                
                for control_id, control_data in failed_controls.items():
                    if control_id not in results['all_controls']:
                        results['all_controls'][control_id] = {
                            'title': control_data['title'],
                            'failed_prod': set(),
                            'failed_uat': set(),
                            'passed_prod': set(),
                            'passed_uat': set()
                        }
                    results['all_controls'][control_id]['failed_uat'].add(server_ip)
    
    return results

def generate_rhel8_audit_report(base_path, output_path):
    """Generate comprehensive RHEL 8 CIS audit report"""
    print("Generating RHEL 8 CIS Audit Report...")
    print(f"Base path: {base_path}")
    
    # Collect results from all servers
    results = collect_server_results(base_path)
    
    # Create Word document
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('RHEL 8 CIS Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # About section
    doc.add_heading('About This Report', 1)
    doc.add_paragraph('This document consolidates CIS (Center for Internet Security) compliance audit results for multiple RHEL 8 systems across Production and UAT environments.')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('CIS Benchmark Reference: ').bold = True
    doc.add_paragraph('This audit follows the official CIS Red Hat Enterprise Linux 8 Benchmark v2.0.0.')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('⚠️ IMPORTANT: Remediation Best Practices').bold = True
    doc.add_paragraph('Test First: Always perform remediation on UAT systems first', style='List Number')
    doc.add_paragraph('Validate: Ensure all applications and services work correctly after remediation', style='List Number')
    doc.add_paragraph('Production Rollout: Only apply changes to Production after successful UAT validation', style='List Number')
    doc.add_paragraph('Golden Image: Once remediation is complete and validated, create a hardened golden image for future deployments', style='List Number')
    doc.add_paragraph('Documentation: Document all changes and maintain an audit trail', style='List Number')
    doc.add_paragraph()
    
    # Servers audited section
    doc.add_heading('Servers Audited', 1)
    
    # Create servers table
    servers_table = doc.add_table(rows=3, cols=2)
    servers_table.style = 'Table Grid'
    
    # Headers
    servers_table.rows[0].cells[0].text = 'Environment'
    servers_table.rows[0].cells[1].text = 'Server IPs'
    for cell in servers_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_color(cell, 'd4edda')
    
    # Production servers
    servers_table.rows[1].cells[0].text = 'Production'
    servers_table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    prod_ips = '\n'.join(sorted(results['prod_servers'])) if results['prod_servers'] else 'None'
    servers_table.rows[1].cells[1].text = prod_ips
    
    # UAT servers
    servers_table.rows[2].cells[0].text = 'UAT'
    servers_table.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True
    uat_ips = '\n'.join(sorted(results['uat_servers'])) if results['uat_servers'] else 'None'
    servers_table.rows[2].cells[1].text = uat_ips
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    total_controls = len(results['all_controls'])
    total_failed = sum(1 for control in results['all_controls'].values() 
                      if control['failed_prod'] or control['failed_uat'])
    total_passed = total_controls - total_failed
    
    doc.add_paragraph(f'Total Controls Evaluated: {total_controls}')
    doc.add_paragraph(f'Passed: {total_passed}')
    doc.add_paragraph(f'Failed: {total_failed}')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('For Remediation: ').bold = True
    p.add_run('Refer to official CIS Benchmark documentation: https://www.cisecurity.org/cis-benchmarks')
    
    doc.add_page_break()
    
    # Detailed findings
    doc.add_heading('Detailed Audit Findings', 1)
    
    if not results['all_controls']:
        doc.add_paragraph('No failed controls found or unable to parse audit results.')
        doc.add_paragraph('Please verify that the RTF files contain valid OpenSCAP output.')
    else:
        # Sort controls by ID
        sorted_controls = sorted(results['all_controls'].keys(), 
                               key=lambda x: [int(p) if p.isdigit() else 0 for p in x.split('.')])
        
        for control_id in sorted_controls:
            control_data = results['all_controls'][control_id]
            cis_details = get_cis_control_details(control_id)
            
            # Control heading
            doc.add_heading(f"{control_id}: {cis_details['title']}", 2)
            
            # Description
            if cis_details.get('description'):
                p = doc.add_paragraph()
                p.add_run('Description: ').bold = True
                doc.add_paragraph(cis_details['description'])
            
            # Impact
            if cis_details.get('impact'):
                p = doc.add_paragraph()
                p.add_run('Impact: ').bold = True
                doc.add_paragraph(cis_details['impact'])
            
            # Remediation
            if cis_details.get('remediation'):
                p = doc.add_paragraph()
                p.add_run('Remediation: ').bold = True
                doc.add_paragraph(cis_details['remediation'])
            
            # Results table
            results_table = doc.add_table(rows=2, cols=2)
            results_table.style = 'Table Grid'
            
            # Headers
            results_table.rows[0].cells[0].text = 'Failed (Production)'
            results_table.rows[0].cells[1].text = 'Failed (UAT)'
            
            for cell in results_table.rows[0].cells:
                cell.paragraphs[0].runs[0].font.bold = True
                set_cell_color(cell, 'f8d7da')
            
            # Failed servers
            failed_prod = '\n'.join(sorted(control_data['failed_prod'])) if control_data['failed_prod'] else 'None'
            failed_uat = '\n'.join(sorted(control_data['failed_uat'])) if control_data['failed_uat'] else 'None'
            
            results_table.rows[1].cells[0].text = failed_prod
            results_table.rows[1].cells[1].text = failed_uat
            
            if control_data['failed_prod'] or control_data['failed_uat']:
                set_cell_color(results_table.rows[1].cells[0], 'f8d7da')
                set_cell_color(results_table.rows[1].cells[1], 'f8d7da')
            
            doc.add_paragraph()
    
    # Save document
    doc.save(output_path)
    print(f'✅ RHEL 8 CIS Audit Report created: {output_path}')
    print(f'Total Controls: {total_controls} | Failed: {total_failed}')

if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python3 rhel8_audit_report_generator.py <base_path>")
        print("Example: python3 rhel8_audit_report_generator.py /Users/satish.korra/Downloads/RHEL-8")
        sys.exit(1)
    
    base_path = sys.argv[1]
    output_path = os.path.join(base_path, f'RHEL-8-CIS-Audit-Report-{datetime.now().strftime("%Y%m%d")}.docx')
    
    generate_rhel8_audit_report(base_path, output_path)