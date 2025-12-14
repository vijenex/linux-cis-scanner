#!/usr/bin/env python3
"""
RHEL-8 UAT CIS Audit Report Generator
Processes RTF files from OpenSCAP scans and generates consolidated Word document
"""
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_color(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)

def extract_ip_from_filename(filename):
    """Extract IP from filename like 172.24.0.136.txt.rtf"""
    match = re.search(r'(\d+\.\d+\.\d+\.\d+)', filename)
    return match.group(1) if match else "N/A"

def parse_rtf_scan_results(rtf_path):
    """Parse OpenSCAP RTF scan results"""
    with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    # Extract IP from filename
    ip_address = extract_ip_from_filename(os.path.basename(rtf_path))
    
    # Extract compliance summary
    passed_match = re.search(r'(\d+)\s+passed', content)
    failed_match = re.search(r'(\d+)\s+failed', content)
    passed = int(passed_match.group(1)) if passed_match else 0
    failed = int(failed_match.group(1)) if failed_match else 0
    
    # Extract failed controls
    failed_controls = []
    fail_pattern = r'\\cf7.*?\\cb3.*?\\strokec7\s+([^\\]+).*?fail'
    for match in re.finditer(fail_pattern, content, re.DOTALL):
        control_title = match.group(1).strip()
        if control_title and len(control_title) > 10:
            failed_controls.append(control_title)
    
    return {
        'ip_address': ip_address,
        'passed': passed,
        'failed': failed,
        'total': passed + failed,
        'failed_controls': failed_controls[:10]  # Limit to first 10
    }

def generate_rhel8_audit_report(rtf_files, output_path):
    """Generate RHEL-8 CIS Audit Report"""
    print(f"Generating RHEL-8 UAT CIS Audit Report...")
    
    # Parse all scan results
    all_results = []
    for rtf_file in rtf_files:
        try:
            result = parse_rtf_scan_results(rtf_file)
            result['filename'] = os.path.basename(rtf_file)
            all_results.append(result)
            print(f"  Processed: {result['ip_address']} - {result['failed']} failed")
        except Exception as e:
            print(f"  Error processing {rtf_file}: {e}")
    
    # Create document
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Red Hat Enterprise Linux 8 CIS Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('UAT Environment - Consolidated Compliance Assessment', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # About This Report
    doc.add_heading('About This Report', 1)
    doc.add_paragraph('This document consolidates CIS (Center for Internet Security) compliance audit results for Red Hat Enterprise Linux 8 systems in the UAT environment.')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('CIS Benchmark Reference: ').bold = True
    p.add_run('This audit follows the official CIS Red Hat Enterprise Linux 8 Benchmark v4.0.0')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Download Official Documentation: ').bold = True
    p.add_run('https://www.cisecurity.org/cis-benchmarks')
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    
    total_systems = len(all_results)
    total_passed = sum(r['passed'] for r in all_results)
    total_failed = sum(r['failed'] for r in all_results)
    total_controls = total_passed + total_failed
    
    doc.add_paragraph(f'Total Systems Audited: {total_systems}')
    doc.add_paragraph(f'Total Controls Evaluated: {total_controls}')
    doc.add_paragraph(f'Total Passed: {total_passed}')
    doc.add_paragraph(f'Total Failed: {total_failed}')
    
    if total_controls > 0:
        compliance_rate = (total_passed / total_controls) * 100
        doc.add_paragraph(f'Overall Compliance Rate: {compliance_rate:.2f}%')
    
    doc.add_page_break()
    
    # System-wise Summary
    doc.add_heading('System-wise Compliance Summary', 1)
    
    # Create summary table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'IP Address'
    header_cells[1].text = 'Total Controls'
    header_cells[2].text = 'Passed'
    header_cells[3].text = 'Failed'
    header_cells[4].text = 'Compliance %'
    
    # Color header
    for cell in header_cells:
        set_cell_color(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Add data rows
    for result in all_results:
        row_cells = table.add_row().cells
        row_cells[0].text = result['ip_address']
        row_cells[1].text = str(result['total'])
        row_cells[2].text = str(result['passed'])
        row_cells[3].text = str(result['failed'])
        
        if result['total'] > 0:
            compliance = (result['passed'] / result['total']) * 100
            row_cells[4].text = f"{compliance:.2f}%"
            
            # Color code compliance
            if compliance >= 90:
                set_cell_color(row_cells[4], 'd4edda')
            elif compliance >= 70:
                set_cell_color(row_cells[4], 'fff3cd')
            else:
                set_cell_color(row_cells[4], 'f8d7da')
    
    doc.add_page_break()
    
    # Detailed Findings per System
    doc.add_heading('Detailed Findings by System', 1)
    
    for result in all_results:
        doc.add_heading(f"System: {result['ip_address']}", 2)
        
        # System summary
        doc.add_paragraph(f"Total Controls: {result['total']}")
        doc.add_paragraph(f"Passed: {result['passed']}")
        doc.add_paragraph(f"Failed: {result['failed']}")
        
        if result['total'] > 0:
            compliance = (result['passed'] / result['total']) * 100
            doc.add_paragraph(f"Compliance Rate: {compliance:.2f}%")
        
        # Failed controls table
        if result['failed_controls']:
            doc.add_paragraph()
            doc.add_paragraph().add_run('Sample Failed Controls:').bold = True
            
            ctrl_table = doc.add_table(rows=1, cols=2)
            ctrl_table.style = 'Table Grid'
            
            # Header
            ctrl_header = ctrl_table.rows[0].cells
            ctrl_header[0].text = '#'
            ctrl_header[1].text = 'Control Title'
            
            for cell in ctrl_header:
                set_cell_color(cell, 'f8d7da')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add failed controls
            for idx, control in enumerate(result['failed_controls'], 1):
                row = ctrl_table.add_row().cells
                row[0].text = str(idx)
                row[1].text = control
        
        doc.add_paragraph()
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
    
    # Remediation Section
    doc.add_page_break()
    doc.add_heading('Remediation Guidance', 1)
    
    doc.add_paragraph('For detailed remediation steps for each failed control, please refer to:')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('1. Official CIS Benchmark: ').bold = True
    p.add_run('CIS Red Hat Enterprise Linux 8 Benchmark v4.0.0')
    
    p = doc.add_paragraph()
    p.add_run('2. Download Link: ').bold = True
    p.add_run('https://www.cisecurity.org/cis-benchmarks')
    
    p = doc.add_paragraph()
    p.add_run('3. OpenSCAP Documentation: ').bold = True
    p.add_run('https://www.open-scap.org/resources/documentation/')
    
    doc.add_paragraph()
    doc.add_paragraph('Each control in the CIS Benchmark includes:')
    doc.add_paragraph('• Detailed description of the security requirement', style='List Bullet')
    doc.add_paragraph('• Step-by-step remediation procedures', style='List Bullet')
    doc.add_paragraph('• Audit commands to verify compliance', style='List Bullet')
    doc.add_paragraph('• Impact assessment of implementing the control', style='List Bullet')
    
    # Save document
    doc.save(output_path)
    print(f'✅ Report created: {output_path}')
    print(f'Systems: {total_systems} | Total Controls: {total_controls} | Compliance: {compliance_rate:.2f}%')

def main():
    # Input folder with RTF files
    input_folder = "/Users/satish.korra/Downloads/RHEL-8/UAT"
    
    # Output path
    output_path = "/Users/satish.korra/Downloads/RHEL-8/RHEL-8-UAT-CIS-Audit-Report.docx"
    
    # Find all RTF files
    rtf_files = []
    for file in os.listdir(input_folder):
        if file.endswith('.rtf'):
            rtf_files.append(os.path.join(input_folder, file))
    
    if not rtf_files:
        print(f"No RTF files found in {input_folder}")
        return
    
    print(f"Found {len(rtf_files)} RTF files to process")
    
    # Generate report
    generate_rhel8_audit_report(rtf_files, output_path)

if __name__ == "__main__":
    main()
