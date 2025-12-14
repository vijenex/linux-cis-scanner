#!/usr/bin/env python3
"""Generate RHEL-8 CIS Audit Report in Windows-style format"""
import os
import csv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from rhel8_parser import parse_cis_benchmark

def set_cell_color(cell, color):
    """Set background color for table cell"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)

def normalize_title(title):
    """Normalize title for better matching"""
    t = title.lower()
    # Normalize common variations
    t = t.replace('ensure ', '').replace('configure ', '').replace('verify ', '')
    t = t.replace(' is ', ' ').replace(' are ', ' ').replace(' must be ', ' ')
    t = t.replace(' option ', ' ').replace(' setting ', ' ')
    t = t.replace('disabled', 'disable').replace('enabled', 'enable')
    t = t.replace('configured', 'configure').replace('installed', 'install')
    return t.strip()

def fuzzy_match(title1, title2):
    """Fuzzy matching for control titles"""
    t1 = normalize_title(title1)
    t2 = normalize_title(title2)
    
    # Direct substring match
    if t1 in t2 or t2 in t1:
        return 0.9
    
    # Extract key words
    ignore = {'the', 'is', 'are', 'a', 'an', 'and', 'or', 'for', 'to', 'of', 'in', 'on', 'be', 'as', 'at', 'by', 'it', 'with', 'from', 'that', 'this', 'have', 'has', 'must', 'should'}
    words1 = set(w for w in t1.split() if w not in ignore and len(w) > 3)
    words2 = set(w for w in t2.split() if w not in ignore and len(w) > 3)
    
    if not words1 or not words2:
        return 0
    
    common = len(words1 & words2)
    return (common / max(len(words1), len(words2))) * 0.8 if common > 0 else 0

def generate_report(scan_results_dir, cis_rtf_path, output_path):
    """Generate RHEL-8 audit report"""
    
    # Parse CIS benchmark
    print("Parsing CIS benchmark...")
    cis_controls = parse_cis_benchmark(cis_rtf_path)
    print(f"Loaded {len(cis_controls)} CIS controls")
    
    # Read scan results
    print("Reading scan results...")
    all_controls = {}
    systems = {}
    
    for file in os.listdir(scan_results_dir):
        if file.endswith('_controls.csv'):
            ip = file.replace('_controls.csv', '')
            systems[ip] = 'nonprod'  # Default to nonprod, can be changed
            
            with open(os.path.join(scan_results_dir, file), 'r') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    title = row['Title']
                    result = row['Result']
                    severity = row['Severity']
                    
                    if title not in all_controls:
                        all_controls[title] = {
                            'title': title,
                            'severity': severity,
                            'passed_prod': set(),
                            'failed_prod': set(),
                            'passed_nonprod': set(),
                            'failed_nonprod': set()
                        }
                    
                    env = systems[ip]
                    if result == 'Pass':
                        all_controls[title][f'passed_{env}'].add(ip)
                    elif result == 'Fail':
                        all_controls[title][f'failed_{env}'].add(ip)
    
    print(f"Found {len(all_controls)} unique controls across {len(systems)} systems")
    
    # Match with CIS controls
    print("Matching with CIS benchmark...")
    matched = 0
    for title in all_controls:
        best_match = None
        best_score = 0
        best_cid = None
        
        for cid, cdata in cis_controls.items():
            score = fuzzy_match(title, cdata['title'])
            if score > best_score:
                best_score = score
                best_match = cdata
                best_cid = cid
        
        if best_match and best_score > 0.2:
            all_controls[title]['cis'] = best_match
            matched += 1
        else:
            all_controls[title]['cis'] = {}
    
    print(f"Matched {matched}/{len(all_controls)} controls with CIS benchmark")
    
    # Generate Word document
    print("Generating Word document...")
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Red Hat Enterprise Linux 8 CIS Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Consolidated Compliance Assessment', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # About section
    doc.add_paragraph('About This Report').bold = True
    doc.add_paragraph('This document consolidates CIS (Center for Internet Security) compliance audit results for multiple Red Hat Enterprise Linux 8 systems across Production and Non-Production environments.')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('CIS Benchmark Reference:').bold = True
    doc.add_paragraph('This audit follows the official CIS Red Hat Enterprise Linux 8 Benchmark v4.0.0. Download the official documentation at: https://www.cisecurity.org/cis-benchmarks')
    doc.add_paragraph()
    
    # Best practices
    p = doc.add_paragraph()
    p.add_run('⚠️ IMPORTANT: Remediation Best Practices').bold = True
    doc.add_paragraph('Test First: Always perform remediation on Non-Production systems first', style='List Number')
    doc.add_paragraph('Validate: Ensure all applications and services work correctly after remediation', style='List Number')
    doc.add_paragraph('Production Rollout: Only apply changes to Production after successful Non-Prod validation', style='List Number')
    doc.add_paragraph('Golden Image: Once remediation is complete and validated, create a hardened golden image for future deployments', style='List Number')
    doc.add_paragraph('Documentation: Document all changes and maintain an audit trail', style='List Number')
    doc.add_paragraph()
    
    # Executive summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(f'Total Controls Evaluated: {len(all_controls)}')
    doc.add_paragraph(f'Systems Scanned: {len(systems)}')
    
    total_pass = sum(1 for c in all_controls.values() if c['passed_nonprod'] or c['passed_prod'])
    total_fail = sum(1 for c in all_controls.values() if c['failed_nonprod'] or c['failed_prod'])
    doc.add_paragraph(f'Overall Pass: {total_pass} | Overall Fail: {total_fail}')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('For Remediation: ').bold = True
    p.add_run('Refer to official CIS Benchmark documentation for detailed remediation steps: https://www.cisecurity.org/cis-benchmarks')
    doc.add_page_break()
    
    # Add controls
    for idx, (title, data) in enumerate(sorted(all_controls.items()), 1):
        cis = data.get('cis', {})
        
        # Control heading
        control_id = cis.get('id', '')
        heading_text = f"{control_id}: {title}" if control_id else f"{idx}. {title}"
        doc.add_heading(heading_text, 2)
        
        # Section
        if control_id:
            section = '.'.join(control_id.split('.')[:2])
            doc.add_paragraph(f"Section: {section}")
        doc.add_paragraph()
        
        # Description
        if cis.get('description'):
            doc.add_paragraph().add_run('Description:').bold = True
            doc.add_paragraph(cis['description'])
        
        # Impact
        if cis.get('impact'):
            doc.add_paragraph().add_run('Impact:').bold = True
            doc.add_paragraph(cis['impact'])
        
        # Remediation
        if cis.get('remediation'):
            doc.add_paragraph().add_run('Remediation:').bold = True
            doc.add_paragraph(cis['remediation'])
        
        # Results table
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['Passed (Prod)', 'Passed (Nonprod)', 'Failed (Prod)', 'Failed (Nonprod)']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            if 'Passed' in header:
                set_cell_color(table.rows[0].cells[i], 'd4edda')
            else:
                set_cell_color(table.rows[0].cells[i], 'f8d7da')
            
            for p in table.rows[0].cells[i].paragraphs:
                for r in p.runs:
                    r.font.bold = True
        
        # Data
        table.rows[1].cells[0].text = '\n'.join(sorted(data['passed_prod'])) or 'None'
        table.rows[1].cells[1].text = '\n'.join(sorted(data['passed_nonprod'])) or 'None'
        table.rows[1].cells[2].text = '\n'.join(sorted(data['failed_prod'])) or 'None'
        table.rows[1].cells[3].text = '\n'.join(sorted(data['failed_nonprod'])) or 'None'
        
        # Evidence placeholder for failed controls
        if data['failed_prod'] or data['failed_nonprod']:
            doc.add_paragraph()
            doc.add_paragraph().add_run('Evidence Screenshot:').bold = True
            doc.add_paragraph('[Screenshot placeholder - Add evidence of failure here]')
        
        doc.add_paragraph()
    
    # Save document
    doc.save(output_path)
    print(f"\n✅ Report generated: {output_path}")
    print(f"Controls: {len(all_controls)} | Systems: {len(systems)}")

if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    scan_results_dir = os.path.join(base_dir, 'scan-results/rhel-8')
    cis_rtf_path = os.path.join(base_dir, 'cis-benchmarks/rhel-8/CIS_Red_Hat_Enterprise_Linux_8_Benchmark_v4.0.0.txt.rtf')
    output_path = os.path.join(base_dir, 'generated-reports/Red-Hat-Enterprise-Linux-8-CIS-Audit-Report.docx')
    
    generate_report(scan_results_dir, cis_rtf_path, output_path)
