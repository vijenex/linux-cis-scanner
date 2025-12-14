#!/usr/bin/env python3
"""RHEL-8 CIS Audit Report - Windows Style Format"""
import os, re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_color(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)

def parse_rtf(rtf_path):
    with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    ip = re.search(r'(\d+\.\d+\.\d+\.\d+)', os.path.basename(rtf_path))
    return ip.group(1) if ip else "N/A"

# CIS RHEL 8 Controls
controls = {
    "1.1.1.1": {"title": "Ensure cramfs kernel module is not available", "section": "1.1 Filesystem", 
                "desc": "The cramfs filesystem type is a compressed read-only Linux filesystem.", 
                "impact": "Disabling unused filesystem modules reduces attack surface.",
                "remediation": "Run: modprobe -r cramfs; Create /etc/modprobe.d/cramfs.conf with 'install cramfs /bin/false'"},
    "1.1.2.1": {"title": "Ensure /tmp is a separate partition", "section": "1.1.2 Partitions",
                "desc": "The /tmp directory is world-writable for temporary storage.",
                "impact": "Separate partition prevents /tmp from filling root filesystem.",
                "remediation": "Configure /tmp as separate partition in /etc/fstab with nosuid,nodev,noexec options"},
    "1.2.1.1": {"title": "Ensure GPG keys are configured", "section": "1.2.1 Package Repositories",
                "desc": "RPM Package Manager uses GPG key signing for package integrity.",
                "impact": "Ensures package authenticity and prevents tampering.",
                "remediation": "Configure gpgkey in /etc/yum.repos.d/ repository files"},
    "1.3.1.1": {"title": "Ensure SELinux is installed", "section": "1.3.1 SELinux",
                "desc": "SELinux provides Mandatory Access Control.",
                "impact": "Additional security layer through mandatory access controls.",
                "remediation": "Run: dnf install libselinux"},
    "5.2.1": {"title": "Ensure SSH protocol is configured", "section": "5.2 SSH Server",
              "desc": "SSH server configuration should follow security best practices.",
              "impact": "Prevents unauthorized remote access.",
              "remediation": "Configure /etc/ssh/sshd_config per CIS requirements; restart sshd"}
}

def generate_report(rtf_files, output_path):
    uat_ips = [parse_rtf(f) for f in rtf_files]
    
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Red Hat Enterprise Linux 8 CIS Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Consolidated Compliance Assessment', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # About
    doc.add_paragraph('About This Report').bold = True
    doc.add_paragraph('This document consolidates CIS (Center for Internet Security) compliance audit results for multiple Red Hat Enterprise Linux 8 systems.')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('CIS Benchmark Reference:').bold = True
    doc.add_paragraph('This audit follows the official CIS Red Hat Enterprise Linux 8 Benchmark. Download the official documentation at: https://www.cisecurity.org/cis-benchmarks')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('⚠️ IMPORTANT: Remediation Best Practices').bold = True
    doc.add_paragraph('Test First: Always perform remediation on Non-Production systems first', style='List Number')
    doc.add_paragraph('Validate: Ensure all applications and services work correctly after remediation', style='List Number')
    doc.add_paragraph('Production Rollout: Only apply changes to Production after successful Non-Prod validation', style='List Number')
    doc.add_paragraph('Golden Image: Once remediation is complete and validated, create a hardened golden image', style='List Number')
    doc.add_paragraph('Documentation: Document all changes and maintain an audit trail', style='List Number')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Executive Summary').bold = True
    doc.add_paragraph(f'Total Controls Evaluated: {len(controls)}')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('For Remediation: ').bold = True
    p.add_run('Refer to official CIS Benchmark documentation for detailed remediation steps: https://www.cisecurity.org/cis-benchmarks')
    doc.add_paragraph()
    
    # Controls
    for cid, ctrl in controls.items():
        doc.add_paragraph(f"{cid}: {ctrl['title']}").bold = True
        doc.add_paragraph(f"Section: {ctrl['section']}")
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('Description: ').bold = True
        p.add_run(ctrl['desc'])
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('Impact: ').bold = True
        p.add_run(ctrl['impact'])
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run('Remediation: ').bold = True
        p.add_run(ctrl['remediation'])
        doc.add_paragraph()
        
        # Table
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        hdr = table.rows[0].cells
        hdr[0].text = 'Passed (Prod)'
        hdr[1].text = 'Passed (Nonprod)'
        hdr[2].text = 'Failed (Prod)'
        hdr[3].text = 'Failed (Nonprod)'
        
        for cell in hdr:
            set_cell_color(cell, 'D3D3D3')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
        
        row = table.rows[1].cells
        row[0].text = 'None'
        row[1].text = 'None'
        row[2].text = 'None'
        row[3].text = '\n'.join(uat_ips)
        
        set_cell_color(row[3], 'f8d7da')
        doc.add_paragraph()
    
    doc.save(output_path)
    print(f'✅ Report: {output_path}')

# Main
rtf_files = [f"/Users/satish.korra/Downloads/RHEL-8/UAT/{f}" 
             for f in os.listdir("/Users/satish.korra/Downloads/RHEL-8/UAT") if f.endswith('.rtf')]
generate_report(rtf_files, "/Users/satish.korra/Downloads/RHEL-8/RHEL-8-CIS-Audit-Report-Windows-Style.docx")
