#!/usr/bin/env python3
"""RHEL-8 CIS Audit Report - Exact Windows Style"""
import os, re, csv
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import PyPDF2

def set_cell_color(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)

def parse_pdf(pdf_path):
    """Parse CIS RHEL-8 PDF"""
    controls = {}
    with open(pdf_path, 'rb') as f:
        pdf = PyPDF2.PdfReader(f)
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
    
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    i = 0
    while i < len(lines):
        line = lines[i]
        match = re.match(r'^(\d+\.\d+(?:\.\d+)*)\s+\(L\d\)', line)
        if match:
            control_id = match.group(1)
            controls[control_id] = {'description': '', 'remediation': '', 'impact': ''}
            i += 1
            current_section = None
            while i < len(lines):
                line = lines[i]
                if re.match(r'^\d+\.\d+(?:\.\d+)*\s+\(L\d\)', line):
                    break
                if 'Description:' in line or 'Rationale:' in line:
                    current_section = 'description'
                elif 'Impact:' in line:
                    current_section = 'impact'
                elif 'Remediation:' in line or 'Audit:' in line:
                    current_section = 'remediation'
                elif 'Default Value:' in line or 'References:' in line:
                    current_section = None
                elif current_section and line:
                    controls[control_id][current_section] += line + ' '
                i += 1
            continue
        i += 1
    return controls

def parse_csv(csv_path):
    findings = []
    with open(csv_path, 'r', encoding='utf-8-sig', errors='ignore') as f:
        for row in csv.DictReader(f):
            findings.append(row)
    return findings

def generate_report(uat_dir, prod_dir, pdf_path, output_path):
    print("Generating RHEL-8 CIS Audit Report...")
    
    # Parse PDF
    print("Parsing CIS PDF...")
    cis_controls = parse_pdf(pdf_path)
    print(f"Parsed {len(cis_controls)} controls")
    
    # Collect results
    control_results = {}
    
    for env, path in [('nonprod', uat_dir), ('prod', prod_dir)]:
        if not os.path.exists(path):
            continue
        for file in os.listdir(path):
            if not file.endswith('_detailed.csv'):
                continue
            csv_path = os.path.join(path, file)
            ip = re.search(r'(\d+\.\d+\.\d+\.\d+)', file)
            if not ip:
                continue
            server = ip.group(1)
            
            for row in parse_csv(csv_path):
                control_id = row.get('Rule_ID', '').strip()
                if not control_id:
                    continue
                
                if control_id not in control_results:
                    control_results[control_id] = {
                        'title': row.get('Title', '').strip(),
                        'passed_prod': set(),
                        'failed_prod': set(),
                        'passed_nonprod': set(),
                        'failed_nonprod': set()
                    }
                
                result = row.get('Result', '').strip().lower()
                if result == 'pass':
                    control_results[control_id][f'passed_{env}'].add(server)
                elif result == 'fail':
                    control_results[control_id][f'failed_{env}'].add(server)
    
    # Create document
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    
    title = doc.add_heading('Red Hat Enterprise Linux 8 CIS Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Consolidated Compliance Assessment', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_paragraph('About This Report').bold = True
    doc.add_paragraph('This document consolidates CIS (Center for Internet Security) compliance audit results for multiple Red Hat Enterprise Linux 8 systems across Production and Non-Production environments.')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('CIS Benchmark Reference:').bold = True
    doc.add_paragraph('This audit follows the official CIS Red Hat Enterprise Linux 8 Benchmark. Download the official documentation at: https://www.cisecurity.org/cis-benchmarks')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('⚠️ IMPORTANT: Remediation Best Practices').bold = True
    doc.add_paragraph('Test First: Always perform remediation on Non-Production systems first', style='List Number')
    doc.add_paragraph('Validate: Ensure all applications and services work correctly after remediation', style='List Number')
    doc.add_paragraph('Production Rollout: Only apply changes to Production after successful Non-Prod validation', style='List Number')
    doc.add_paragraph('Golden Image: Once remediation is complete and validated, create a hardened golden image for future deployments', style='List Number')
    doc.add_paragraph('Documentation: Document all changes and maintain an audit trail', style='List Number')
    doc.add_paragraph()
    
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(f'Total Controls Evaluated: {len(control_results)}')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('For Remediation: ').bold = True
    p.add_run('Refer to official CIS Benchmark documentation for detailed remediation steps: https://www.cisecurity.org/cis-benchmarks')
    doc.add_page_break()
    
    # Add controls
    control_ids = sorted(control_results.keys(), key=lambda x: [int(p) if p.isdigit() else 0 for p in x.split('.')])
    
    for control_id in control_ids:
        data = control_results[control_id]
        
        doc.add_heading(f"{control_id}: {data['title']}", 2)
        doc.add_paragraph(f"Section: {control_id.split('.')[0]}")
        
        cis_data = cis_controls.get(control_id, {})
        if cis_data.get('description'):
            doc.add_paragraph().add_run('Description:').bold = True
            doc.add_paragraph(cis_data['description'].strip())
        
        if cis_data.get('impact'):
            doc.add_paragraph().add_run('Impact:').bold = True
            doc.add_paragraph(cis_data['impact'].strip())
        
        if cis_data.get('remediation'):
            doc.add_paragraph().add_run('Remediation:').bold = True
            doc.add_paragraph(cis_data['remediation'].strip())
        
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Passed (Prod)'
        table.rows[0].cells[1].text = 'Passed (Nonprod)'
        table.rows[0].cells[2].text = 'Failed (Prod)'
        table.rows[0].cells[3].text = 'Failed (Nonprod)'
        
        set_cell_color(table.rows[0].cells[0], 'd4edda')
        set_cell_color(table.rows[0].cells[1], 'd4edda')
        set_cell_color(table.rows[0].cells[2], 'f8d7da')
        set_cell_color(table.rows[0].cells[3], 'f8d7da')
        
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
        
        table.rows[1].cells[0].text = '\n'.join(sorted(data['passed_prod'])) or 'None'
        table.rows[1].cells[1].text = '\n'.join(sorted(data['passed_nonprod'])) or 'None'
        table.rows[1].cells[2].text = '\n'.join(sorted(data['failed_prod'])) or 'None'
        table.rows[1].cells[3].text = '\n'.join(sorted(data['failed_nonprod'])) or 'None'
        
        if data['failed_prod'] or data['failed_nonprod']:
            doc.add_paragraph()
            doc.add_paragraph().add_run('Evidence Screenshot:').bold = True
            doc.add_paragraph('[Screenshot placeholder - Add evidence of failure here]')
        
        doc.add_paragraph()
    
    doc.save(output_path)
    print(f'✅ Report: {output_path}')
    print(f'Controls: {len(control_results)}')

# Main
uat_dir = "/Users/satish.korra/Downloads/RHEL-8/UAT"
prod_dir = "/Users/satish.korra/Downloads/RHEL-8/Prod"
pdf_path = "/Users/satish.korra/Downloads/CIS_Red_Hat_Enterprise_Linux_8_Benchmark_v4.0.0.pdf"
output_path = "/Users/satish.korra/Downloads/RHEL-8/RHEL-8-CIS-Audit-Report.docx"

generate_report(uat_dir, prod_dir, pdf_path, output_path)
