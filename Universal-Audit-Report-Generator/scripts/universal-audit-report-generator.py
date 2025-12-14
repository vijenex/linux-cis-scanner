#!/usr/bin/env python3
"""
Universal Windows Server CIS Audit Report Generator
Supports: Windows Server 2016, 2019, 2022, 2025
"""
import os
import csv
import sys
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from striprtf.striprtf import rtf_to_text
import re

def set_cell_color(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)

def parse_rtf(rtf_path):
    """Parse CIS Benchmark RTF file"""
    if not os.path.exists(rtf_path):
        print(f"Warning: RTF file not found: {rtf_path}")
        return {}
    
    with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
        text = rtf_to_text(f.read())
    
    controls = {}
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    
    i = 0
    while i < len(lines):
        line = lines[i]
        match = re.match(r'^(\d+\.\d+(?:\.\d+)*(?:\.\d+)?)\s+\(L\d\)', line)
        if match:
            control_id = match.group(1)
            controls[control_id] = {'description': '', 'remediation': '', 'impact': ''}
            i += 1
            
            current_section = None
            while i < len(lines):
                line = lines[i]
                if re.match(r'^\d+\.\d+(?:\.\d+)*(?:\.\d+)?\s+\(L\d\)', line):
                    break
                
                if 'Profile Applicability:' in line or 'Description:' in line or 'Rationale:' in line:
                    current_section = 'description'
                elif 'Impact:' in line:
                    current_section = 'impact'
                elif 'Audit:' in line or 'Remediation:' in line:
                    current_section = 'remediation'
                elif 'Default Value:' in line or 'References:' in line or 'CIS Controls:' in line:
                    current_section = None
                elif current_section and line:
                    controls[control_id][current_section] += line + ' '
                
                i += 1
            continue
        i += 1
    
    return controls

def parse_csv(csv_path):
    """Parse scan results CSV"""
    findings = []
    with open(csv_path, 'r', encoding='utf-8-sig') as f:
        for row in csv.DictReader(f):
            findings.append(row)
    return findings

def generate_report(scan_dir, rtf_path, output_path, version):
    """Generate CIS Audit Report"""
    print(f"Generating Windows Server {version} CIS Audit Report...")
    print(f"Scan directory: {scan_dir}")
    print(f"RTF file: {rtf_path}")
    
    # Parse RTF
    print("Parsing RTF...")
    cis_controls = parse_rtf(rtf_path)
    print(f"Parsed {len(cis_controls)} controls from RTF")
    
    # Collect results with Prod/Nonprod separation
    control_results = {}
    prod_patterns = ['172.18.', '172.20.', '172.22.', '172.23.', '172.27.']
    nonprod_exception = '172.21.1.229'
    
    # Check both Prod and Nonprod directories
    base_path = os.path.dirname(scan_dir)
    version_num = version
    prod_dir = os.path.join(base_path, f'Prod-{version_num}')
    nonprod_dir = os.path.join(base_path, f'Nonprod-{version_num}')
    
    for env, path in [('prod', prod_dir), ('nonprod', nonprod_dir)]:
        if not os.path.exists(path):
            continue
        for server_dir in os.listdir(path):
            csv_file = os.path.join(path, server_dir, 'vijenex-cis-results.csv')
            if not os.path.exists(csv_file):
                continue
            
            for row in parse_csv(csv_file):
                control_id = row['Id'].strip()
                title = row['Title'].strip()
                
                # Skip DC only controls
                if '(DC only)' in title or '(DC Only)' in title:
                    continue
                
                if control_id not in control_results:
                    control_results[control_id] = {
                        'title': title,
                        'section': row['Section'].strip(),
                        'passed_prod': set(),
                        'failed_prod': set(),
                        'passed_nonprod': set(),
                        'failed_nonprod': set()
                    }
                
                if row['Status'] == 'PASS':
                    control_results[control_id][f'passed_{env}'].add(server_dir)
                elif row['Status'] == 'FAIL':
                    control_results[control_id][f'failed_{env}'].add(server_dir)
    
    # Create document
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading(f'Windows Server {version} CIS Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # About
    doc.add_heading('About This Report', 1)
    doc.add_paragraph(f'This document consolidates CIS (Center for Internet Security) compliance audit results for multiple Windows Server {version} systems across Production and Non-Production environments.')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('CIS Benchmark Reference:').bold = True
    doc.add_paragraph(f'This audit follows the official CIS Microsoft Windows Server {version} Benchmark.')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('⚠️ IMPORTANT: Remediation Best Practices').bold = True
    doc.add_paragraph('Test First: Always perform remediation on Non-Production systems first', style='List Number')
    doc.add_paragraph('Validate: Ensure all applications and services work correctly after remediation', style='List Number')
    doc.add_paragraph('Production Rollout: Only apply changes to Production after successful Non-Prod validation', style='List Number')
    doc.add_paragraph('Golden Image: Once remediation is complete and validated, create a hardened golden image for future deployments', style='List Number')
    doc.add_paragraph('Documentation: Document all changes and maintain an audit trail', style='List Number')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Excluded Controls:').bold = True
    doc.add_paragraph('This audit excludes the following control categories as they are not applicable to this environment:')
    doc.add_paragraph('Domain Controller Controls: All controls marked "(DC only)" have been excluded', style='List Bullet')
    doc.add_paragraph('Windows 10/11 Controls: Controls specific to workstations have been excluded', style='List Bullet')
    doc.add_paragraph('MSS (Legacy) Controls: Deprecated Microsoft Security Settings have been excluded', style='List Bullet')
    doc.add_paragraph()
    
    # Servers Audited
    doc.add_heading('Servers Audited', 1)
    prod_servers = set()
    nonprod_servers = set()
    for data in control_results.values():
        prod_servers.update(data['passed_prod'])
        prod_servers.update(data['failed_prod'])
        nonprod_servers.update(data['passed_nonprod'])
        nonprod_servers.update(data['failed_nonprod'])
    
    servers_table = doc.add_table(rows=2, cols=2)
    servers_table.style = 'Table Grid'
    servers_table.rows[0].cells[0].text = 'Prod'
    servers_table.rows[0].cells[1].text = 'Nonprod'
    for cell in servers_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    servers_table.rows[1].cells[0].text = '\n'.join(sorted(prod_servers)) or 'None'
    servers_table.rows[1].cells[1].text = '\n'.join(sorted(nonprod_servers)) or 'None'
    doc.add_paragraph()
    
    # Executive Summary
    total_controls = len(control_results)
    # Get first server CSV for summary
    first_csv = None
    if os.path.exists(prod_dir):
        for sd in os.listdir(prod_dir):
            csv_file = os.path.join(prod_dir, sd, 'vijenex-cis-results.csv')
            if os.path.exists(csv_file):
                first_csv = csv_file
                break
    if not first_csv and os.path.exists(nonprod_dir):
        for sd in os.listdir(nonprod_dir):
            csv_file = os.path.join(nonprod_dir, sd, 'vijenex-cis-results.csv')
            if os.path.exists(csv_file):
                first_csv = csv_file
                break
    
    if first_csv:
        rows = parse_csv(first_csv)
        passed = sum(1 for r in rows if r['Status'] == 'PASS' and '(DC only)' not in r.get('Title', ''))
        failed = sum(1 for r in rows if r['Status'] == 'FAIL' and '(DC only)' not in r.get('Title', ''))
    else:
        passed = failed = 0
    
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(f'Total Controls Evaluated: {total_controls}')
    doc.add_paragraph(f'Passed: {passed}')
    doc.add_paragraph(f'Failed: {failed}')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('For Remediation: ').bold = True
    p.add_run('Refer to official CIS Benchmark documentation: https://www.cisecurity.org/cis-benchmarks')
    doc.add_page_break()
    
    # Sort controls
    control_ids = sorted(control_results.keys(), 
                        key=lambda x: [int(p) if p.isdigit() else 0 for p in x.split('.')])
    
    # Add each control
    missing_desc = 0
    for control_id in control_ids:
        data = control_results[control_id]
        
        # Skip DC only controls
        if '(DC only)' in data['title'] or '(DC Only)' in data['title']:
            continue
        
        # Map shorthand IDs to full IDs (for 2025)
        display_id = control_id
        if control_id.startswith('57.3.'):
            display_id = '18.10.' + control_id
        elif control_id.startswith('43.6.'):
            display_id = '18.10.' + control_id
        
        doc.add_heading(f"{display_id}: {data['title']}", 2)
        doc.add_paragraph(f"Section: {data['section']}")
        
        # Add CIS details
        cis_data = cis_controls.get(display_id, {})
        if cis_data.get('description') and cis_data['description'].strip():
            doc.add_paragraph().add_run('Description:').bold = True
            doc.add_paragraph(cis_data['description'].strip())
        else:
            missing_desc += 1
            
        if cis_data.get('impact') and cis_data['impact'].strip():
            doc.add_paragraph().add_run('Impact:').bold = True
            doc.add_paragraph(cis_data['impact'].strip())
            
        if cis_data.get('remediation') and cis_data['remediation'].strip():
            doc.add_paragraph().add_run('Remediation:').bold = True
            doc.add_paragraph(cis_data['remediation'].strip())
        
        # Create table with 4 columns
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Passed (Prod)'
        table.rows[0].cells[1].text = 'Passed (Nonprod)'
        table.rows[0].cells[2].text = 'Failed (Prod)'
        table.rows[0].cells[3].text = 'Failed (Nonprod)'
        
        set_cell_color(table.rows[0].cells[0], 'd4edda')
        set_cell_color(table.rows[0].cells[1], 'd4edda')
        set_cell_color(table.rows[0].cells[2], 'f8d7da')
        set_cell_color(table.rows[0].cells[3], 'f8d7da')
        
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        table.rows[1].cells[0].text = '\n'.join(sorted(list(data['passed_prod']))) or 'None'
        table.rows[1].cells[1].text = '\n'.join(sorted(list(data['passed_nonprod']))) or 'None'
        table.rows[1].cells[2].text = '\n'.join(sorted(list(data['failed_prod']))) or 'None'
        table.rows[1].cells[3].text = '\n'.join(sorted(list(data['failed_nonprod']))) or 'None'
        
        set_cell_color(table.rows[1].cells[0], 'ffffff')
        set_cell_color(table.rows[1].cells[1], 'ffffff')
        set_cell_color(table.rows[1].cells[2], 'f8d7da')
        set_cell_color(table.rows[1].cells[3], 'f8d7da')
        
        # Add Evidence Screenshot for failed controls
        if data['failed_prod'] or data['failed_nonprod']:
            doc.add_paragraph()
            doc.add_paragraph().add_run('Evidence Screenshot:').bold = True
            doc.add_paragraph('[Screenshot placeholder - Add evidence of failure here]')
        
        doc.add_paragraph()
    
    doc.save(output_path)
    print(f'✅ Report created: {output_path}')
    print(f'Controls: {total_controls} | Passed: {passed} | Failed: {failed}')
    print(f'RTF Controls: {len(cis_controls)} | Missing descriptions: {missing_desc}')

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("Usage: python3 universal-audit-report-generator.py <scan_dir> <rtf_path> <version>")
        print("Example: python3 universal-audit-report-generator.py /path/to/Prod-2025 /path/to/benchmark.rtf 2025")
        sys.exit(1)
    
    scan_dir = sys.argv[1]
    rtf_path = sys.argv[2]
    version = sys.argv[3]
    
    output_path = os.path.join(scan_dir, f'Windows-Server-{version}-CIS-Audit-Report.docx')
    generate_report(scan_dir, rtf_path, output_path, version)
